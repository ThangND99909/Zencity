import re
import json
import os
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from sentence_transformers import SentenceTransformer
import numpy as np

# Thư viện đọc tài liệu
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# OCR hỗ trợ fallback nếu PDF là ảnh
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class ZenPreprocessor:
    """
    Lớp tiền xử lý (preprocessor) cho dữ liệu ZenCity Foundation:
    - Làm sạch văn bản
    - Phân loại tài liệu giáo dục
    - Chunk văn bản
    - Sinh embeddings
    - Lưu JSONL/NPZ
    """

    def __init__(self, model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"):
        """
        Khởi tạo mô hình SentenceTransformer (đa ngôn ngữ: Anh – Việt)
        """
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)

        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(levelname)s - %(message)s"
        )
        logging.info(f"✅ ZenPreprocessor initialized with model: {model_name}")

    # ============================================================
    # 1️⃣ Làm sạch văn bản
    # ============================================================
    def clean_text(self, text: str) -> str:
        """Loại bỏ HTML, script, ký tự đặc biệt, footer, sitemap,..."""
        if not text:
            return ""

        # Xóa script/style + HTML
        text = re.sub(r'<(script|style).*?>.*?</\1>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<.*?>', '', text)

        # Loại bỏ footer, banner, copyright, sitemap
        noise_patterns = [
            r'Follow us.*', r'Subscribe.*', r'Contact us.*',
            r'©.*\d{4}.*', r'Terms of Use.*', r'Privacy Policy.*',
            r'All rights reserved.*', r'Sitemap.*', r'Search.*'
        ]
        for pat in noise_patterns:
            text = re.sub(pat, '', text, flags=re.IGNORECASE)

        # Thay ký tự HTML đặc biệt
        html_entities = {
            '&nbsp;': ' ', '&amp;': '&', '&quot;': '"', '&apos;': "'",
            '&lt;': '<', '&gt;': '>', '\u2013': '-', '\u2014': '-',
            '\u2022': '•', '\u00a0': ' '
        }
        for k, v in html_entities.items():
            text = text.replace(k, v)

        # Chuẩn hóa khoảng trắng
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    # ============================================================
    # 2️⃣ Phân loại tài liệu ZenCity
    # ============================================================
    def process_documents(self, documents):
        """Làm sạch + phân loại tài liệu theo nội dung giáo dục ZenCity"""
        processed_docs = []

        for doc in documents:
            content = self.clean_text(doc.page_content)
            if not content:
                continue

            metadata = doc.metadata or {}
            doc_type = "general"

            # Phân loại theo từ khóa giáo dục
            if re.search(r"\b(course|class|lesson|chương trình|khóa học|bài học)\b", content, re.IGNORECASE):
                doc_type = "course_material"
            elif re.search(r"\b(teacher|giáo viên|training|đào tạo)\b", content, re.IGNORECASE):
                doc_type = "teacher_training"
            elif re.search(r"\b(student|học viên|trẻ em|learner|học sinh)\b", content, re.IGNORECASE):
                doc_type = "student_info"
            elif re.search(r"\b(event|sự kiện|workshop|seminar)\b", content, re.IGNORECASE):
                doc_type = "event_info"
            elif re.search(r"\b(blog|news|tin tức|bài viết)\b", content, re.IGNORECASE):
                doc_type = "blog_post"
            elif re.search(r"\b(contact|liên hệ|hotline|email|địa chỉ)\b", content, re.IGNORECASE):
                doc_type = "contact_info"

            processed_doc = Document(
                page_content=content,
                metadata={**metadata, "type": doc_type}
            )
            processed_docs.append(processed_doc)

        logging.info(f"🧹 Processed {len(processed_docs)} documents.")
        return processed_docs

    # ============================================================
    # 3️⃣ Chia chunk văn bản
    # ============================================================
    def split_documents(self, documents, chunk_size=1000, chunk_overlap=200):
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", " "],
            length_function=len,
        )
        chunks = splitter.split_documents(documents)
        logging.info(f"📄 Split into {len(chunks)} chunks.")
        return chunks

    # ============================================================
    # 4️⃣ Sinh vector embeddings
    # ============================================================
    def embed_documents(self, documents):
        texts = [doc.page_content for doc in documents if doc.page_content.strip()]
        embeddings = self.model.encode(
            texts,
            show_progress_bar=True,
            batch_size=32,
            normalize_embeddings=True
        )
        logging.info(f"🧠 Created {len(embeddings)} embeddings.")
        return embeddings

    # ============================================================
    # 5️⃣ Lưu dữ liệu
    # ============================================================
    def save_to_jsonl(self, documents, filename):
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        with open(filename, "w", encoding="utf-8") as f:
            for doc in documents:
                f.write(json.dumps({
                    "content": doc.page_content,
                    "metadata": doc.metadata
                }, ensure_ascii=False) + "\n")
        logging.info(f"💾 Saved {len(documents)} docs → {filename}")

    def save_embeddings(self, embeddings, documents, out_path="data/zencity_embeddings.npz"):
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        np.savez_compressed(
            out_path,
            embeddings=embeddings,
            metadata=[doc.metadata for doc in documents]
        )
        logging.info(f"💾 Saved embeddings → {out_path}")

    # ============================================================
    # 6️⃣ Pipeline chính: clean → classify → chunk
    # ============================================================
    def clean_and_chunk(self, raw_docs, chunk_size=1000, chunk_overlap=200):
        processed = self.process_documents(raw_docs)
        chunks = self.split_documents(processed, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return chunks

    # ============================================================
    # 7️⃣ Đọc PDF (text + OCR fallback)
    # ============================================================
    def read_pdf(self, file_path: str) -> str:
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            logging.warning(f"PDF read error: {e}")

        if not text.strip() and OCR_AVAILABLE:
            try:
                poppler_path = r"E:\Poppler\poppler-24.07.0\Library\bin"  # Update nếu cần
                images = convert_from_path(file_path, poppler_path=poppler_path)
                for img in images:
                    text += pytesseract.image_to_string(img, lang="vie+eng") + "\n"
            except Exception as e:
                logging.error(f"OCR failed for {file_path}: {e}")

        return self.clean_text(text)

    # ============================================================
    # 8️⃣ Đọc TXT / DOCX
    # ============================================================
    def read_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return self.clean_text(f.read())

    def read_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return self.clean_text(text)
